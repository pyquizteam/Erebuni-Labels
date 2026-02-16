import streamlit as st
import pandas as pd
import io
import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont 

EXCEL_FILE = '–ë–æ—á–∫–∏ 3-—è —Ñ—É—Ä–∞.xlsx'
WORD_OUTPUT = 'Labels_150x100.docx'
PDF_OUTPUT = 'Labels_150x100.pdf'

LABEL_W_MM = 150
LABEL_H_MM = 100
LABEL_W, LABEL_H = LABEL_W_MM * mm, LABEL_H_MM * mm

def setup_pdf_fonts():
    """Registers fonts. If Arial.ttf isn't in your folder, it uses Helvetica."""
    try:
        pdfmetrics.registerFont(TTFont('ArialCustom', 'Arial.ttf'))
        pdfmetrics.registerFont(TTFont('ArialBoldCustom', 'Arial-Bold.ttf'))
        return 'ArialCustom', 'ArialBoldCustom'
    except:
        return 'Helvetica', 'Helvetica-Bold'



def create_word_file(data):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(LABEL_W_MM), Mm(LABEL_H_MM)
    section.top_margin, section.bottom_margin = Mm(10), Mm(10)
    section.left_margin, section.right_margin = Mm(12), Mm(12)

    for i, row in data.iterrows():
        if i > 0: doc.add_page_break()
        table = doc.add_table(rows=1, cols=3)
        table.width = Mm(LABEL_W_MM - 24)
        
        mid = table.cell(0, 1).paragraphs[0]
        mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = mid.add_run("–ü–ò–¶–¶–ê –°–û–£–°")
        r.bold, r.font.size = True, Pt(8)

        right = table.cell(0, 2).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = right.add_run("(–ü–†–û–î–£–ö–¢ –ü–ê–°–¢–ï–†–ò–ó–û–í–ê–ù)")
        r.bold, r.font.size = True, Pt(6)

        def tight_p(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run(text)
            run.bold, run.font.size = True, Pt(11)

        p_date = pd.to_datetime(row['–î–∞—Ç–∞ –ü—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞']).strftime('%d.%m.%Y')
        e_date = pd.to_datetime(row['–ì–æ–¥–µ–Ω –¥–æ']).strftime('%d.%m.%Y')
        ph = f"{float(row['PH']):.2f}"

        tight_p(f"–ù–æ–º–µ—Ä –ø–∞—Ä—Ç–∏–∏: {row['–ù–æ–º–µ—Ä –ü–∞—Ä—Ç–∏–∏']}  /  BRIX: {row['BRIX']}%  /  PH: {ph}")
        tight_p(f"–í–µ—Å –ù–µ—Ç—Ç–æ: {row['–ù–µ—Ç—Ç–æ —Å–æ—É—Å–∞']}  /  –í–µ—Å –ë—Ä—É—Ç—Ç–æ: {row['–ë—Ä—É—Ç—Ç–æ –±–æ—á–µ–∫']}")
        tight_p(f"–ò–∑–≥–æ—Ç–æ–≤–ª–µ–Ω–æ: {p_date}  /  –ì–æ–¥–µ–Ω –¥–æ: {e_date}")

        legal = doc.add_paragraph()
        legal.paragraph_format.line_spacing = 1.3
        run1 = legal.add_run(
            "--------------------------------------------------------------------------------\n"
            "–°–æ—Å—Ç–∞–≤: –¢–æ–º–∞—Ç—ã, —Å–æ–ª—å –ø–∏—â–µ–≤–∞—è, —Ä–µ–≥—É–ª—è—Ç–æ—Ä –∫–∏—Å–ª–æ—Ç–Ω–æ—Å—Ç–∏ (–ª–∏–º–æ–Ω–Ω–∞—è –∫–∏—Å–ª–æ—Ç–∞).\n"
            "–ü–∏—â–µ–≤–∞—è —Ü–µ–Ω–Ω–æ—Å—Ç—å –Ω–∞ 100 –≥. –ø—Ä–æ–¥—É–∫—Ç–∞: —É–≥–ª–µ–≤–æ–¥—ã - 9,0–≥, –∂–∏—Ä—ã - 0,0–≥\n"
            "—ç–Ω–µ—Ä–≥–µ—Ç–∏—á–µ—Å–∫–∞—è —Ü–µ–Ω–Ω–æ—Å—Ç—å - 36,0 –∫–∫–∞–ª/153,0 –ö–î–∂.\n"
            "–•—Ä–∞–Ω–∏—Ç—å –≤ —Å—É—Ö–æ–º, –ø—Ä–æ—Ö–ª–∞–¥–Ω–æ–º –º–µ—Å—Ç–µ –ø—Ä–∏ —Ç–µ–º–ø–µ—Ä–∞—Ç—É—Ä–µ –æ—Ç +0¬∞–° –¥–æ +25¬∞C –∏\n"
            "–æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ–π –≤–ª–∞–∂–Ω–æ—Å—Ç–∏ –≤–æ–∑–¥—É—Ö–∞ –Ω–µ –±–æ–ª–µ–µ 75%. –ü–æ—Å–ª–µ –≤—Å–∫—Ä—ã—Ç–∏—è —Ö—Ä–∞–Ω–∏—Ç—å –ø—Ä–∏\n"
            "—Ç–µ–º–ø–µ—Ä–∞—Ç—É—Ä–µ –æ—Ç +0¬∞–° –¥–æ +5¬∞–° –Ω–µ –±–æ–ª–µ–µ 24 —á–∞—Å–∞.\n"
            "–°—Ä–æ–∫ –≥–æ–¥–Ω–æ—Å—Ç–∏: 12 –º–µ—Å—è—Ü–µ–≤ —Å –¥–∞—Ç—ã –∏–∑–≥–æ—Ç–æ–≤–ª–µ–Ω–∏—è.\n"
            '–ò–∑–≥–æ—Ç–æ–≤–∏—Ç–µ–ª—å: –û–û–û "–≠—Ä–µ–±—É–Ω–∏ –¢—Ä–µ–π–¥ –ì—Ä—É–ø".\n'
            "–Æ—Ä–∏–¥–∏—á–µ—Å–∫–∏–π –∞–¥—Ä–µ—Å: –†–ê, –≥. –ï—Ä–µ–≤–∞–Ω, —É–ª. –ê–≤–µ—Ç–∏—Å—è–Ω, –¥–æ–º 23.\n"
            "–ê–¥—Ä–µ—Å –ø—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞: –†–µ—Å–ø—É–±–ª–∏–∫–∞ –ê—Ä–º–µ–Ω–∏—è, –ê—Ä–∞—Ä–∞—Ç—Å–∫–∏–π —Ä–µ–≥., –≥. –ê—Ä—Ç–∞—à–∞—Ç,\n"
            "—É–ª. –°–∞–º–≤–µ–ª–∞ –ê–∫–æ–ø—è–Ω–∞, —Å—Ç—Ä. 173.\n"
            "–¢–µ–ª: (+374) 77-733-388 Email: erebunit@gmail.com\n"
            "–¢–£ AM 51192101. 9183-2023\n"
        )
        run1.font.size = Pt(9)

        run2 = legal.add_run("–ü—Ä–æ–∏–∑–≤–µ–¥–µ–Ω–æ –≤ –†–µ—Å–ø—É–±–ª–∏–∫–µ –ê—Ä–º–µ–Ω–∏—è.")
        run2.bold = True
        run2.font.size = Pt(9)

    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()


def create_pdf_file(data):
    f_reg, f_bold = setup_pdf_fonts()
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=(LABEL_W, LABEL_H))
    
    MARGIN_X, MARGIN_TOP = 5 * mm, 7 * mm
    pallet_counter = 1

    for index, row in data.iterrows():
        c.setFont(f_bold, 12)
        c.drawCentredString(LABEL_W / 2, LABEL_H - MARGIN_TOP, "–ü–ò–¶–¶–ê –°–û–£–°")
        
        c.setFont(f_bold, 10)
        c.drawRightString(LABEL_W - MARGIN_X, LABEL_H - MARGIN_TOP, "(–ü–†–û–î–£–ö–¢ –ü–ê–°–¢–ï–†–ò–ó–û–í–ê–ù)")

        p_date = pd.to_datetime(row['–î–∞—Ç–∞ –ü—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞']).strftime('%d.%m.%Y')
        e_date = pd.to_datetime(row['–ì–æ–¥–µ–Ω –¥–æ']).strftime('%d.%m.%Y')
        ph = f"{float(row['PH']):.2f}"

        y = LABEL_H - MARGIN_TOP - 5 * mm
        c.setFont(f_bold, 10)

        c.drawString(MARGIN_X, y,
            f"–ù–æ–º–µ—Ä –ø–∞—Ä—Ç–∏–∏: {row['–ù–æ–º–µ—Ä –ü–∞—Ä—Ç–∏–∏']}  /  BRIX: {row['BRIX']}%  /  PH: {ph}")
        y -= 4 * mm
        c.drawString(MARGIN_X, y,
            f"–í–µ—Å –ù–µ—Ç—Ç–æ: {row['–ù–µ—Ç—Ç–æ —Å–æ—É—Å–∞']}  /  –í–µ—Å –ë—Ä—É—Ç—Ç–æ: {row['–ë—Ä—É—Ç—Ç–æ –±–æ—á–µ–∫']}")
        y -= 4 * mm
        c.drawString(MARGIN_X, y,
            f"–ò–∑–≥–æ—Ç–æ–≤–ª–µ–Ω–æ: {p_date}  /  –ì–æ–¥–µ–Ω –¥–æ: {e_date}")

        c.line(MARGIN_X, y - 3 * mm, LABEL_W - MARGIN_X, y - 3 * mm)

        c.setFont(f_reg, 10)
        y -= 10 * mm

        text = c.beginText(MARGIN_X, y)
        text.setLeading(13)  

        for line in [
            "–°–æ—Å—Ç–∞–≤: –¢–æ–º–∞—Ç—ã, —Å–æ–ª—å –ø–∏—â–µ–≤–∞—è, —Ä–µ–≥—É–ª—è—Ç–æ—Ä –∫–∏—Å–ª–æ—Ç–Ω–æ—Å—Ç–∏ (–ª–∏–º–æ–Ω–Ω–∞—è –∫–∏—Å–ª–æ—Ç–∞).",
            "–ü–∏—â–µ–≤–∞—è —Ü–µ–Ω–Ω–æ—Å—Ç—å –Ω–∞ 100 –≥. –ø—Ä–æ–¥—É–∫—Ç–∞: —É–≥–ª–µ–≤–æ–¥—ã - 9,0–≥, –∂–∏—Ä—ã - 0,0–≥",
            "—ç–Ω–µ—Ä–≥–µ—Ç–∏—á–µ—Å–∫–∞—è —Ü–µ–Ω–Ω–æ—Å—Ç—å - 36,0 –∫–∫–∞–ª/153,0 –ö–î–∂.",
            "–•—Ä–∞–Ω–∏—Ç—å –≤ —Å—É—Ö–æ–º, –ø—Ä–æ—Ö–ª–∞–¥–Ω–æ–º –º–µ—Å—Ç–µ –ø—Ä–∏ —Ç–µ–º–ø–µ—Ä–∞—Ç—É—Ä–µ –æ—Ç +0¬∞–° –¥–æ +25¬∞C –∏",
            "–æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ–π –≤–ª–∞–∂–Ω–æ—Å—Ç–∏ –≤–æ–∑–¥—É—Ö–∞ –Ω–µ –±–æ–ª–µ–µ 75%. –ü–æ—Å–ª–µ –≤—Å–∫—Ä—ã—Ç–∏—è —Ö—Ä–∞–Ω–∏—Ç—å –ø—Ä–∏",
            "—Ç–µ–º–ø–µ—Ä–∞—Ç—É—Ä–µ –æ—Ç +0¬∞–° –¥–æ +5¬∞–° –Ω–µ –±–æ–ª–µ–µ 24 —á–∞—Å–∞.",
            "–°—Ä–æ–∫ –≥–æ–¥–Ω–æ—Å—Ç–∏: 12 –º–µ—Å—è—Ü–µ–≤ —Å –¥–∞—Ç—ã –∏–∑–≥–æ—Ç–æ–≤–ª–µ–Ω–∏—è.",
            '–ò–∑–≥–æ—Ç–æ–≤–∏—Ç–µ–ª—å: –û–û–û "–≠—Ä–µ–±—É–Ω–∏ –¢—Ä–µ–π–¥ –ì—Ä—É–ø".',
            "–Æ—Ä–∏–¥–∏—á–µ—Å–∫–∏–π –∞–¥—Ä–µ—Å: –†–ê, –≥. –ï—Ä–µ–≤–∞–Ω, —É–ª. –ê–≤–µ—Ç–∏—Å—è–Ω, –¥–æ–º 23.",
            "–ê–¥—Ä–µ—Å –ø—Ä–æ–∏–∑–≤–æ–¥—Å—Ç–≤–∞: –†–µ—Å–ø—É–±–ª–∏–∫–∞ –ê—Ä–º–µ–Ω–∏—è, –ê—Ä–∞—Ä–∞—Ç—Å–∫–∏–π —Ä–µ–≥., –≥. –ê—Ä—Ç–∞—à–∞—Ç,",
            "—É–ª. –°–∞–º–≤–µ–ª–∞ –ê–∫–æ–ø—è–Ω–∞, —Å—Ç—Ä. 173.",
            "–¢–µ–ª: (+374) 77-733-388 Email: erebunit@gmail.com",
            "",
            "–¢–£ AM 51192101. 9183-2023",
        ]:
            text.textLine(line)

        c.drawText(text)

        c.setFont(f_bold, 10)
        c.drawString(MARGIN_X, text.getY() - 3, "–ü—Ä–æ–∏–∑–≤–µ–¥–µ–Ω–æ –≤ –†–µ—Å–ø—É–±–ª–∏–∫–µ –ê—Ä–º–µ–Ω–∏—è.")

        LOGO_SIZE = 10 * mm
        logo_gap = 1 * mm
        LOGO_Y = 3 * mm

        if os.path.exists("logo_right.png"):
            c.drawImage("logo_right.png",
                        LABEL_W - LOGO_SIZE - MARGIN_X,
                        LOGO_Y,
                        width=LOGO_SIZE, height=LOGO_SIZE,
                        preserveAspectRatio=True, mask='auto')

        if os.path.exists("logo_left.png"):
            c.drawImage("logo_left.png",
                        LABEL_W - 2*LOGO_SIZE - MARGIN_X - logo_gap,
                        LOGO_Y,
                        width=LOGO_SIZE, height=LOGO_SIZE,
                        preserveAspectRatio=True, mask='auto')

        c.showPage()

        if (index + 1) % 4 == 0:
            c.setFont(f_bold, 28)
            c.drawCentredString(LABEL_W / 2, 75 * mm, f"–ü–ê–õ–ï–¢–ê \u2116 {pallet_counter:02d}")
            c.setFont(f_bold, 24)
            val_net = row['–ù–µ—Ç—Ç–æ —Å–æ—É—Å–∞ –Ω–∞ –ø–∞–ª–ª–µ—Ç–µ']
            val_gross = row['–ë—Ä—É—Ç—Ç–æ –ø–∞–ª–ª–µ—Ç–∞']
            
            net_str = f"{float(val_net):.1f}".replace('.', ',') if pd.notna(val_net) else "0,0"
            gross_str = f"{float(val_gross):.1f}".replace('.', ',') if pd.notna(val_gross) else "0,0"
            
            LEFT_ALIGN_X = 25 * mm 
            c.drawString(LEFT_ALIGN_X, 45 * mm, f"–í–µ—Å –ù–µ—Ç—Ç–æ  -  {net_str}")
            c.drawString(LEFT_ALIGN_X, 20 * mm, f"–í–µ—Å –ë—Ä—É—Ç—Ç–æ  -  {gross_str}")
            pallet_counter += 1
            c.showPage()

    c.save()
    return buffer.getvalue()

st.set_page_config(page_title="Erebuni Label Gen", page_icon="üì¶")
st.title("Label Generator")
st.info("Upload your Excel file (–ë–æ—á–∫–∏) to generate labels.")

uploaded_file = st.file_uploader("Choose Excel File", type=['xlsx'])

if uploaded_file:
    df = pd.read_excel(uploaded_file, skiprows=4)
    df.columns = [" ".join(str(c).split()) for c in df.columns]
    for col in ['–ù–µ—Ç—Ç–æ —Å–æ—É—Å–∞ –Ω–∞ –ø–∞–ª–ª–µ—Ç–µ', '–ë—Ä—É—Ç—Ç–æ –ø–∞–ª–ª–µ—Ç–∞']:
        if col in df.columns:
            df[col] = df[col].ffill()
    df = df[df['–ù–æ–º–µ—Ä –ü–∞—Ä—Ç–∏–∏'].notna()].copy()
    df = df.reset_index(drop=True)
    st.success(f"Loaded {len(df)} labels from Excel.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Generate PDF"):
            pdf_data = create_pdf_file(df)
            st.download_button("Download PDF", pdf_data, "Labels.pdf", "application/pdf")
            
    with col2:
        if st.button("Generate Word"):
            word_data = create_word_file(df)
            st.download_button("Download Word", word_data, "Labels.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
